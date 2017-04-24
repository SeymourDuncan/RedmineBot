from docx import Document
from docx.shared import Cm, Pt
from consts import Paths, Reports
from mytypes import DocumentFile

import re

head_names = ['№ п/п', 'Объект тестирования', 'Сценарий выполнения', 'Ожидаемый результат', 'Инициатор', 'Итог тестирования']
# cm
column_width = [1, 5, 7, 6, 3, 5]


# Заполнение данными строку таблицы
# row - строка
# data - list c данными
# isHead - признак заголовка таблицы. Влияет на формат
def fillRowData(row, data, isHead = False):

    def styleRun(run):
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(10)
        if isHead:
            run.bold = True

    for i, text in enumerate(data):
        char_idx = 0
        pgph = row[i].add_paragraph()
        # ищем картинки
        matches = re.finditer((r"!!") + '(?P<content>[^\>]*?)' + (r"!!"), text)
        # ходим по картинкам
        for m in matches:
            # добавляем текст до картинки
            piece = text[char_idx: m.start()]
            if piece:
                run = pgph.add_run(piece)
                styleRun(run)
            # добавляем картинку
            pgph.add_run().add_picture(m.group('content'), width=Cm(column_width[i] - 0.5))
            char_idx = m.end()

        # добавляем то что осталось после картинок
        piece = text[char_idx:]
        run = pgph.add_run(piece)
        styleRun(run)
        row[i].width = Cm(column_width[i]).emu

def BuildDocx(data, version = ''):
    document = Document(Paths.temp_file)

    # заголовок
    paragraph = document.add_paragraph()
    run = paragraph.add_run(Reports.header_tmp.format(version))
    run.style = 'Заголовок 1 Знак'
    paragraph.paragraph_format.first_line_indent = Cm(3.0)

    # цикл по категориям
    idx = 1
    for ctg, issues in data.items():
        # вписываем категорию
        paragraph = document.add_paragraph('Таблица {} {}'.format(idx, ctg))
        # paragraph.paragraph_format.space_after = Pt(12)
        paragraph.style = 'Caption'

        # строим таблицу
        table = document.add_table(rows=len(issues) + 1, cols=6)
        table.autofit = False

        table.style = 'Table Grid'
        row_id = 0
        row = table.rows[row_id].cells
        # заголовок
        fillRowData(row, head_names, True)
        # данные
        for iss in issues:
            row_id += 1
            row = table.rows[row_id].cells
            iss.insert(0, '{}.'.format(row_id))
            fillRowData(row, iss)
        idx+=1

    document.save(Paths.test_prot_file)
    return DocumentFile(Paths.test_prot_file)


# BuildDocx(data)

